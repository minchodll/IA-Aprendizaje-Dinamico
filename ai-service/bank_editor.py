"""Plantilla Word para que un docente proponga un tema nuevo (con sus
ejercicios) para el banco semilla, y el parser/validador que convierte esa
plantilla llena de vuelta a la forma que espera bank_loader.py.

Flujo: GET /plantilla-banco (descarga) -> el docente la llena en Word ->
POST /banco/preview (sube el .docx, se valida sin guardar nada) ->
POST /banco/confirmar (guarda lo ya validado en banco-ejercicios.json).

Importante: esto NO reentrena los modelos automáticamente. Agregar
ejercicios malos sin revisar podría degradar el clasificador, así que
guardar y reentrenar quedan como pasos separados y deliberados.
"""

import io
import json

from docx import Document
from docx.shared import RGBColor

from bank_loader import BANK_PATH

CATEGORIAS_EXISTENTES = [
    "excel", "word", "powerpoint", "internet_redes",
    "sistema_operativo", "hardware_software", "seguridad_informatica",
]
TIPOS_VALIDOS = {"opcion_multiple", "respuesta_corta"}
DIFICULTADES_VALIDAS = {"basico", "intermedio", "avanzado"}

CAMPOS_TEMA = ["Nombre del Tema", "Categoría", "Competencia", "Conceptos Clave"]
CAMPOS_EJERCICIO = [
    "Enunciado", "Tipo", "Opciones", "Respuesta Correcta", "Dificultad",
    "Retroalimentación si acierta", "Retroalimentación si falla", "Respuestas alternativas",
]

PRIMARY = RGBColor(0x02, 0x80, 0x90)
MUTED = RGBColor(0x5C, 0x7A, 0x7A)


def _tabla_clave_valor(doc, filas):
    tabla = doc.add_table(rows=0, cols=2)
    tabla.style = "Table Grid"
    for etiqueta, valor in filas:
        fila = tabla.add_row()
        fila.cells[0].text = etiqueta
        fila.cells[0].paragraphs[0].runs[0].bold = True
        fila.cells[1].text = valor
    return tabla


def _ejercicio_filas(enunciado="", tipo="", opciones="", respuesta_correcta="",
                      dificultad="", retro_correcta="", retro_incorrecta="", alternativas=""):
    return list(zip(CAMPOS_EJERCICIO, [
        enunciado, tipo, opciones, respuesta_correcta, dificultad,
        retro_correcta, retro_incorrecta, alternativas,
    ]))


def generar_plantilla() -> bytes:
    doc = Document()

    titulo = doc.add_heading("Plantilla: Nuevo Tema para el Banco de Ejercicios", level=1)
    titulo.runs[0].font.color.rgb = PRIMARY

    doc.add_paragraph(
        "Llena esta plantilla para proponer un tema nuevo (o agregar más ejercicios a uno "
        "existente) para el sistema de generación automática de exámenes. No cambies el texto "
        "de la columna izquierda de las tablas — solo escribe en la columna derecha."
    )
    p = doc.add_paragraph()
    p.add_run("Categorías que ya existen: ").bold = True
    p.add_run(", ".join(CATEGORIAS_EXISTENTES) + ". ")
    p.add_run(
        "Si usas una de estas, tus ejercicios se agregan a ese tema. Si escribes una categoría "
        "nueva, se crea un tema nuevo — pero la IA no podrá clasificarlo bien hasta que tenga "
        "varios ejercicios curados y se reentrene el modelo."
    )
    p2 = doc.add_paragraph()
    p2.add_run("Para cada ejercicio: ").bold = True
    p2.add_run(
        "Tipo debe ser exactamente \"opcion_multiple\" o \"respuesta_corta\". Dificultad debe ser "
        "exactamente \"basico\", \"intermedio\" o \"avanzado\". En Opciones y Respuestas "
        "alternativas, separa cada una con el símbolo | (barra vertical)."
    )

    doc.add_heading("Datos del Tema", level=2)
    _tabla_clave_valor(doc, [
        ("Nombre del Tema", "Ej: Microsoft Excel"),
        ("Categoría", "Ej: excel"),
        ("Competencia", "Ej: Manejo de hojas de cálculo: celdas, fórmulas y funciones básicas"),
        ("Conceptos Clave", "Ej: celdas, fórmulas básicas, funciones SUMA/PROMEDIO"),
    ])

    doc.add_heading("Ejercicios", level=2)
    doc.add_paragraph(
        "Hay 2 ejemplos ya llenos (bórralos o edítalos) y 6 bloques vacíos. Si necesitas más de "
        "8 ejercicios, selecciona una tabla de ejercicio completa (con su título) y copiapégala "
        "las veces que necesites."
    )

    ejemplos = [
        (
            "Ejercicio de ejemplo 1 (opción múltiple)",
            _ejercicio_filas(
                enunciado="¿Qué es una celda en Excel?",
                tipo="opcion_multiple",
                opciones="La intersección entre una fila y una columna|Un tipo de gráfico|El nombre de un archivo de Excel|Una fórmula matemática",
                respuesta_correcta="La intersección entre una fila y una columna",
                dificultad="basico",
                retro_correcta="Correcto. Cada celda se identifica por su columna (letra) y fila (número), por ejemplo A1.",
                retro_incorrecta="Repasa el concepto de celda: es donde se cruzan una fila y una columna.",
                alternativas="(no aplica para opción múltiple)",
            ),
        ),
        (
            "Ejercicio de ejemplo 2 (respuesta corta)",
            _ejercicio_filas(
                enunciado="Explica con tus palabras qué hace la función PROMEDIO en Excel.",
                tipo="respuesta_corta",
                opciones="(no aplica para respuesta corta)",
                respuesta_correcta="Calcula el promedio (media aritmética) de un conjunto de valores numéricos seleccionados.",
                dificultad="intermedio",
                retro_correcta="Bien explicado: PROMEDIO calcula la media aritmética de los valores seleccionados.",
                retro_incorrecta="Tu respuesta no menciona que PROMEDIO calcula la media aritmética; revisa ese concepto.",
                alternativas="Suma los valores de un rango y los divide entre la cantidad de datos|Saca la media de varios números en las celdas seleccionadas",
            ),
        ),
    ]
    for titulo_bloque, filas in ejemplos:
        h = doc.add_heading(titulo_bloque, level=3)
        h.runs[0].font.color.rgb = MUTED
        _tabla_clave_valor(doc, filas)

    for i in range(3, 9):
        h = doc.add_heading(f"Ejercicio {i}", level=3)
        h.runs[0].font.color.rgb = MUTED
        _tabla_clave_valor(doc, _ejercicio_filas())

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _valor_por_etiqueta(tabla, etiquetas_buscadas):
    """Busca en una tabla clave/valor la fila cuya primera celda empieza con
    alguna de las etiquetas dadas (normalizado), y devuelve el texto de la
    segunda celda. Tolera que el docente no borre los acentos o mayúsculas."""
    for fila in tabla.rows:
        if len(fila.cells) < 2:
            continue
        etiqueta = fila.cells[0].text.strip().lower()
        for buscada in etiquetas_buscadas:
            if etiqueta.startswith(buscada):
                return fila.cells[1].text.strip()
    return ""


def parsear_docx(contenido: bytes) -> dict:
    doc = Document(io.BytesIO(contenido))
    tablas = doc.tables
    errores = []
    advertencias = []

    if len(tablas) == 0:
        return {"tema": None, "errores": ["El documento no tiene ninguna tabla — usa la plantilla sin modificar su estructura."], "advertencias": []}

    datos_tema = tablas[0]
    nombre_tema = _valor_por_etiqueta(datos_tema, ["nombre del tema"])
    categoria = _valor_por_etiqueta(datos_tema, ["categoría", "categoria"]).strip().lower()
    competencia = _valor_por_etiqueta(datos_tema, ["competencia"])
    conceptos_raw = _valor_por_etiqueta(datos_tema, ["conceptos clave"])
    conceptos_clave = [c.strip() for c in conceptos_raw.split(",") if c.strip()]

    if not nombre_tema:
        errores.append("Falta el 'Nombre del Tema' en la tabla de Datos del Tema.")
    if not categoria:
        errores.append("Falta la 'Categoría' en la tabla de Datos del Tema.")
    elif categoria not in CATEGORIAS_EXISTENTES:
        advertencias.append(
            f"'{categoria}' no es una de las 7 categorías existentes ({', '.join(CATEGORIAS_EXISTENTES)}). "
            "Se creará como categoría nueva, pero la IA no la reconocerá bien hasta reentrenar con varios ejercicios."
        )
    if not competencia:
        advertencias.append("No escribiste la 'Competencia' del tema.")
    if not conceptos_clave:
        advertencias.append("No escribiste 'Conceptos Clave' del tema.")

    ejercicios = []
    for idx, tabla in enumerate(tablas[1:], start=1):
        enunciado = _valor_por_etiqueta(tabla, ["enunciado"])
        if not enunciado or enunciado.startswith("(no aplica"):
            continue  # bloque vacio sin usar, se ignora silenciosamente

        tipo = _valor_por_etiqueta(tabla, ["tipo"]).strip().lower()
        opciones_raw = _valor_por_etiqueta(tabla, ["opciones"])
        respuesta_correcta = _valor_por_etiqueta(tabla, ["respuesta correcta"])
        dificultad = _valor_por_etiqueta(tabla, ["dificultad"]).strip().lower()
        retro_correcta = _valor_por_etiqueta(tabla, ["retroalimentación si acierta", "retroalimentacion si acierta"])
        retro_incorrecta = _valor_por_etiqueta(tabla, ["retroalimentación si falla", "retroalimentacion si falla"])
        alternativas_raw = _valor_por_etiqueta(tabla, ["respuestas alternativas"])

        prefijo = f"Ejercicio #{idx} (\"{enunciado[:40]}...\"):" if len(enunciado) > 40 else f"Ejercicio #{idx} (\"{enunciado}\"):"

        if tipo not in TIPOS_VALIDOS:
            errores.append(f"{prefijo} Tipo debe ser 'opcion_multiple' o 'respuesta_corta', no '{tipo}'.")
            continue
        if dificultad not in DIFICULTADES_VALIDAS:
            errores.append(f"{prefijo} Dificultad debe ser 'basico', 'intermedio' o 'avanzado', no '{dificultad}'.")
            continue
        if not respuesta_correcta:
            errores.append(f"{prefijo} Falta la Respuesta Correcta.")
            continue

        ejercicio = {
            "enunciado": enunciado,
            "tipo": tipo,
            "respuesta_correcta": respuesta_correcta,
            "dificultad": dificultad,
            "retroalimentacion_correcta": retro_correcta or "Correcto.",
            "retroalimentacion_incorrecta": retro_incorrecta or "Incorrecto, revisa el tema.",
        }

        if tipo == "opcion_multiple":
            opciones = [o.strip() for o in opciones_raw.split("|") if o.strip() and not o.strip().startswith("(no aplica")]
            if len(opciones) < 2:
                errores.append(f"{prefijo} Necesita al menos 2 Opciones separadas por '|'.")
                continue
            if respuesta_correcta not in opciones:
                advertencias.append(f"{prefijo} La Respuesta Correcta no coincide exactamente con ninguna Opción (revisa mayúsculas/espacios).")
            ejercicio["opciones"] = opciones
        else:
            alternativas = [a.strip() for a in alternativas_raw.split("|") if a.strip() and not a.strip().startswith("(no aplica")]
            if not alternativas:
                advertencias.append(f"{prefijo} No tiene Respuestas alternativas — ayudan a que la IA reconozca más parafraseos.")
            ejercicio["respuestas_alternativas"] = alternativas

        ejercicios.append(ejercicio)

    if not ejercicios and not errores:
        errores.append("No se encontró ningún ejercicio lleno en el documento (todos los bloques están vacíos o son los ejemplos).")

    tema = {
        "categoria": categoria,
        "nombre_tema": nombre_tema,
        "competencia": competencia,
        "conceptos_clave": conceptos_clave,
        "ejercicios": ejercicios,
    }
    return {"tema": tema, "errores": errores, "advertencias": advertencias}


def agregar_tema_al_banco(tema: dict) -> dict:
    with open(BANK_PATH, encoding="utf-8") as f:
        data = json.load(f)

    categoria = tema["categoria"]
    existente = next((t for t in data["temas"] if t["categoria"] == categoria), None)

    if existente is None:
        existente = {
            "categoria": categoria,
            "nombre_tema": tema["nombre_tema"],
            "competencia": tema["competencia"],
            "conceptos_clave": tema["conceptos_clave"],
            "ejercicios": [],
        }
        data["temas"].append(existente)
    else:
        # el tema ya existe: se agregan los ejercicios nuevos, sin pisar los datos del tema
        if tema.get("competencia") and not existente.get("competencia"):
            existente["competencia"] = tema["competencia"]
        for concepto in tema.get("conceptos_clave", []):
            if concepto not in existente["conceptos_clave"]:
                existente["conceptos_clave"].append(concepto)

    indices_existentes = [
        int(e["id"].rsplit("-", 1)[-1])
        for e in existente["ejercicios"]
        if e["id"].rsplit("-", 1)[-1].isdigit()
    ]
    siguiente = (max(indices_existentes) + 1) if indices_existentes else 1

    for ejercicio in tema["ejercicios"]:
        ejercicio["id"] = f"{categoria}-{siguiente}"
        siguiente += 1
        existente["ejercicios"].append(ejercicio)

    with open(BANK_PATH, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    return {
        "categoria": categoria,
        "nombre_tema": existente["nombre_tema"],
        "ejercicios_agregados": len(tema["ejercicios"]),
        "total_ejercicios_en_categoria": len(existente["ejercicios"]),
        "es_categoria_nueva": categoria not in CATEGORIAS_EXISTENTES,
    }
